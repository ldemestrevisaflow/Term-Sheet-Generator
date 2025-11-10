#!/usr/bin/env python3
"""
Generate Term Sheet DOCX by filling in the master PwC template with JSON questionnaire data.
Handles split runs and row-based replacements.

Usage: python generate_term_sheet.py <path_to_json_file>
"""

import json
import sys
import os
import glob
import re
from datetime import datetime
from docx import Document

def find_template():
    """Find the master Term Sheet template file"""
    print("Searching for Term Sheet template...")
    
    # Try exact names first
    possible_names = [
        'Term Sheet - Share Sale - Binding_option1(ID 2740).docx',
        'Term Sheet - Share Sale - Binding_option[ID 2740].docx',
        'Term Sheet - Share Sale - Binding.docx',
        'Term_Sheet_Master.docx',
        'templates/Term Sheet - Share Sale - Binding_option1(ID 2740).docx',
    ]
    
    for template_name in possible_names:
        if os.path.exists(template_name):
            print(f"✓ Found template: {template_name}")
            return template_name
    
    # Fallback: search for any docx file with "Term Sheet" in name
    print("Exact match not found, searching for any 'Term Sheet' docx files...")
    for pattern in ['*Term Sheet*.docx', 'templates/*Term Sheet*.docx']:
        matches = glob.glob(pattern)
        if matches:
            template_file = matches[0]
            print(f"✓ Found template via glob: {template_file}")
            return template_file
    
    # List all files for debugging
    print("\n⚠️  Template not found! Available DOCX files:")
    for docx_file in glob.glob('**/*.docx', recursive=True):
        print(f"  - {docx_file}")
    
    raise FileNotFoundError(
        "Master Term Sheet template not found. "
        "Expected to find a file with 'Term Sheet' in the name ending in .docx"
    )

def get_cell_text(cell):
    """Get all text from a cell's paragraphs"""
    return ''.join([para.text for para in cell.paragraphs])

def set_cell_text(cell, new_text):
    """Set text in a cell by modifying the first paragraph's runs"""
    if cell.paragraphs:
        para = cell.paragraphs[0]
        # Clear all runs in the paragraph
        for run in list(para.runs):
            r = run._element
            r.getparent().remove(r)
        # Add new run with the text
        para.add_run(new_text)

def format_date_word(date_str):
    """Convert date from YYYY-MM-DD to 'DD Month YYYY' format"""
    if not date_str or date_str.strip() == '':
        return ''
    
    try:
        from datetime import datetime
        # Parse the date - handle ISO format
        if 'T' in date_str:  # ISO format like 2025-12-15T00:00:00
            date_obj = datetime.fromisoformat(date_str.split('T')[0])
        else:  # Simple YYYY-MM-DD format
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        
        # Format as "20 December 2025"
        formatted = date_obj.strftime('%d %B %Y')
        # Remove leading zero from day (e.g., "03" -> "3")
        if formatted[0] == '0':
            formatted = formatted[1:]
        return formatted
    except Exception as e:
        print(f"Warning: Could not format date '{date_str}': {e}")
        return date_str

def replace_text_in_runs(runs, replacements):
    """Replace text across multiple runs in a paragraph"""
    if not runs:
        return
    
    # Get full text
    full_text = ''.join([run.text for run in runs])
    
    # Check if any replacement is needed
    needs_replacement = any(key in full_text for key in replacements.keys())
    
    if not needs_replacement:
        return
    
    # Do replacements on full text
    for key, value in replacements.items():
        full_text = full_text.replace(key, str(value))
    
    # Clear runs and set new text
    for run in runs:
        run.text = ""
    
    if runs:
        runs[0].text = full_text

def replace_in_document_general(doc, replacements):
    """Replace placeholders in paragraphs and general document text"""
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_runs(paragraph.runs, replacements)
    
    # Replace in table cells (general)
    for table_idx, table in enumerate(doc.tables):
        # Skip table 1 (party details table - handle separately)
        if table_idx == 1:
            continue
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_runs(paragraph.runs, replacements)
    
    # Replace in headers/footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_text_in_runs(paragraph.runs, replacements)
        for paragraph in section.footer.paragraphs:
            replace_text_in_runs(paragraph.runs, replacements)

def fix_first_page_parties(doc, seller, buyer):
    """Fix the first page table with party names"""
    print("Fixing first page party details...")
    
    if len(doc.tables) > 0:
        table = doc.tables[0]
        
        # Find and replace in Table 0 (the cover page table)
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell_text = get_cell_text(cell)
                
                # Replace party 1 (Seller)
                if '[Insert Party 1 Name]' in cell_text:
                    set_cell_text(cell, f"{seller.get('name', '')} ABN ({seller.get('abn', '')}) (Seller)")
                    print(f"  Updated Seller in table: {seller.get('name', '')}")
                
                # Replace party 2 (Buyer)
                if '[Insert Party 2 Name]' in cell_text:
                    set_cell_text(cell, f"{buyer.get('name', '')} ABN ({buyer.get('abn', '')}) (Buyer)")
                    print(f"  Updated Buyer in table: {buyer.get('name', '')}")

def fix_binding_text(doc, is_binding):
    """Fix the [non-]Binding text based on binding/non-binding selection"""
    print(f"Fixing binding/non-binding text (is_binding={is_binding})...")
    
    for para in doc.paragraphs:
        # Fix Table of Contents or cover page
        if '[non-]' in para.text or 'Non-]' in para.text:
            for run in para.runs:
                if '[non-]' in run.text:
                    if is_binding:
                        run.text = run.text.replace('[non-]', '')
                    else:
                        run.text = run.text.replace('[non-]', 'non-')
                    print(f"  Fixed: {run.text}")
                elif 'Non-]' in run.text:
                    if is_binding:
                        run.text = run.text.replace('Non-]', '')
                    else:
                        run.text = run.text.replace('Non-]', 'Non-')

def fix_recital_a(doc, target):
    """Fix Recital A with target company name"""
    print("Fixing Recital A with target company name...")
    
    for para in doc.paragraphs:
        if '[insert name and ABN of company]' in para.text:
            for run in para.runs:
                if '[insert name and ABN of company]' in run.text:
                    replacement = f"{target.get('name', '')} (ABN {target.get('abn', '')})"
                    run.text = run.text.replace('[insert name and ABN of company]', replacement)
                    print(f"  Updated Recital A: {replacement}")

def fix_signature_blocks(doc, buyer, seller):
    """Fix signature blocks with buyer and seller names"""
    print("Fixing signature blocks...")
    
    signature_block_count = 0
    for para in doc.paragraphs:
        if 'SIGNED by' in para.text and '[INSERT PARTY NAME]' in para.text:
            signature_block_count += 1
            for run in para.runs:
                if '[INSERT PARTY NAME]' in run.text:
                    # First signature block = Buyer, Second = Seller
                    if signature_block_count == 1:
                        run.text = run.text.replace('[INSERT PARTY NAME]', buyer.get('name', ''))
                        print(f"  Buyer signature block: {buyer.get('name', '')}")
                    elif signature_block_count == 2:
                        run.text = run.text.replace('[INSERT PARTY NAME]', seller.get('name', ''))
                        print(f"  Seller signature block: {seller.get('name', '')}")

def fill_party_details_table(table, seller, buyer, target):
    """Fill in the party details table (Table 1) with seller and buyer info"""
    
    print("Filling party details table...")
    
    # Row 3-8: Seller (Party 1)
    # Row 10-15: Buyer (Party 2)  
    # Row 17-22: Escrow Agent (Party 3)
    
    # Seller Name (Row 3, Cell 1)
    try:
        cell_text = get_cell_text(table.rows[3].cells[1])
        if '[insert Party' in cell_text:
            set_cell_text(table.rows[3].cells[1], seller.get('name', ''))
            print(f"  Seller name: {seller.get('name', '')}")
    except:
        pass
    
    # Seller ABN (Row 4, Cell 1)
    try:
        cell_text = get_cell_text(table.rows[4].cells[1])
        if '[insert' in cell_text:
            set_cell_text(table.rows[4].cells[1], seller.get('abn', ''))
            print(f"  Seller ABN: {seller.get('abn', '')}")
    except:
        pass
    
    # Buyer Name (Row 10, Cell 1)
    try:
        cell_text = get_cell_text(table.rows[10].cells[1])
        if '[insert Party' in cell_text:
            set_cell_text(table.rows[10].cells[1], buyer.get('name', ''))
            print(f"  Buyer name: {buyer.get('name', '')}")
    except:
        pass
    
    # Buyer ABN (Row 11, Cell 1)
    try:
        cell_text = get_cell_text(table.rows[11].cells[1])
        if '[insert' in cell_text:
            set_cell_text(table.rows[11].cells[1], buyer.get('abn', ''))
            print(f"  Buyer ABN: {buyer.get('abn', '')}")
    except:
        pass
    
    # Target/Company Name (Party 3 - Row 17, Cell 1)
    try:
        cell_text = get_cell_text(table.rows[17].cells[1])
        if '[insert Party' in cell_text:
            set_cell_text(table.rows[17].cells[1], target.get('name', ''))
            print(f"  Target/Company name: {target.get('name', '')}")
    except:
        pass
    
    # Target ABN (Row 18, Cell 1)
    try:
        cell_text = get_cell_text(table.rows[18].cells[1])
        if '[insert' in cell_text:
            set_cell_text(table.rows[18].cells[1], target.get('abn', ''))
            print(f"  Target ABN: {target.get('abn', '')}")
    except:
        pass

def replace_signature_blocks(doc, buyer, seller):
    """Replace party names in signature blocks"""
    
    print("Filling signature blocks...")
    
    # Find and replace signature block text
    for para_idx, para in enumerate(doc.paragraphs):
        if 'SIGNED by' in para.text and '[INSERT PARTY NAME]' in para.text:
            # Replace with buyer on first occurrence, seller on second
            text = para.text
            if 'SIGNED by [INSERT PARTY NAME]' in text:
                # This is tricky - we need to track which signature block this is
                # For now, replace all with buyer (you may need manual adjustment)
                for run in para.runs:
                    if '[INSERT PARTY NAME]' in run.text:
                        run.text = run.text.replace('[INSERT PARTY NAME]', buyer.get('name', ''))
                        print(f"  Signature block: {buyer.get('name', '')}")

def format_currency(value):
    """Format value as Australian currency"""
    try:
        if value:
            num = float(value)
            return f"A${num:,.2f}"
    except (ValueError, TypeError):
        pass
    return ""

def format_date(value):
    """Format date value - convert to word format like '20 December 2025'"""
    if value:
        return format_date_word(value)
    return ""

def generate_term_sheet(json_file):
    """Generate a Term Sheet DOCX by filling the master template with JSON data"""
    
    print(f"Loading JSON from: {json_file}")
    
    # Read JSON file
    with open(json_file, 'r') as f:
        data = json.load(f)
    
    # Find and load template
    template_file = find_template()
    print(f"Loading template: {template_file}")
    doc = Document(template_file)
    
    # Extract data
    seller = data.get('seller', {})
    buyer = data.get('buyer', {})
    target = data.get('targetCompany', {})
    transaction = data.get('transaction', {})
    dates = data.get('dates', {})
    conditions = data.get('conditions', {})
    warranties = data.get('warranties', {})
    commercial = data.get('commercial', {})
    management = data.get('management', {})
    legal = data.get('legal', {})
    
    # General replacements for placeholders
    replacements = {
        '[Insert Date]': format_date(dates.get('termSheetDate', '')),
        '[insert date]': format_date(dates.get('completionDate', '')),
        '[Insert Amount]': format_currency(transaction.get('purchasePrice')),
        '[insert amount]': format_currency(transaction.get('depositAmount')),
        '[insert number]': str(conditions.get('infoRequestDays', '10')),
        '[30]': str(conditions.get('accessPeriodDays', '30')),
        '[insert address]': target.get('name', ''),
        '[three]': str(commercial.get('nonCompetePeriod', '3')),
        '[12]': str(commercial.get('nonSolicitationPeriod', '12')),
        '[six months]': '6 months',
        '[five]': '5',
        '[insert relevant name]': management.get('directorsResign', ''),
        '[New South Wales]': legal.get('jurisdiction', 'New South Wales'),
        '[Insert list of Suppliers]': legal.get('suppliersSchedule', ''),
        '[Insert list of Customers]': legal.get('customersSchedule', ''),
        '[insert details]': '',
        '[insert details of representations]': '',
    }
    
    print("Replacing general placeholders...")
    replace_in_document_general(doc, replacements)
    
    print("Applying specific fixes...")
    # Get binding status
    is_binding = legal.get('termSheetType', 'binding') == 'binding'
    
    # Apply all fixes
    fix_first_page_parties(doc, seller, buyer)
    fix_binding_text(doc, is_binding)
    fix_recital_a(doc, target)
    
    print("Filling party details table...")
    # Handle party details table specially (Table 1)
    if len(doc.tables) > 1:
        fill_party_details_table(doc.tables[1], seller, buyer, target)
    
    print("Fixing signature blocks...")
    fix_signature_blocks(doc, buyer, seller)
    
    # Save document
    output_dir = 'output'
    os.makedirs(output_dir, exist_ok=True)
    output_file = os.path.join(output_dir, 'generated_TermSheet.docx')
    
    print(f"Saving to: {output_file}")
    doc.save(output_file)
    
    print(f"✅ Term Sheet generated successfully!")
    print(f"   File: {output_file}")
    print(f"   Size: {os.path.getsize(output_file)} bytes")
    
    return output_file

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python generate_term_sheet.py <json_file>")
        sys.exit(1)
    
    json_file = sys.argv[1]
    
    if not os.path.exists(json_file):
        print(f"Error: File not found: {json_file}")
        sys.exit(1)
    
    try:
        generate_term_sheet(json_file)
        print("Success!")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
